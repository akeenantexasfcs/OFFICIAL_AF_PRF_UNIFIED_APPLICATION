"""
Unified PRF + AF Portfolio Optimizer — Word Report Generator.

Generates a branded TFC Word document from Joint Optimizer results.
Adapted from af_report_generator.py to support mixed PRF + AF portfolios.
"""
import numpy as np
import pandas as pd
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker

from prf_constants import INTERVAL_ORDER_11
from af_constants import SEASON_LABELS, get_buyup_intervals, get_cat_interval, interval_to_months

# =============================================================================
# TFC BRAND COLORS (matches PRF report_generator.py exactly)
# =============================================================================
_CLR_PRIMARY_GREEN = RGBColor(0x5E, 0x97, 0x32)
_CLR_SLATE_BLUE = RGBColor(0x5B, 0x70, 0x7F)
_CLR_MUTED_ROSE = RGBColor(0x9D, 0x5F, 0x58)
_CLR_DARK_FOREST_HEX = '2D3A2E'
_CLR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_CLR_BLACK = RGBColor(0x00, 0x00, 0x00)
_WARM_CREAM = '#F5F1E8'

# Chart colors for Joint vs Independent
_HEX_GREEN = '#5E9732'
_HEX_SLATE = '#5B707F'
_HEX_ROSE = '#9D5F58'

METRIC_DISPLAY_NAMES = {
    'sharpe': 'Risk-Adjusted (Sharpe)',
    'sortino': 'Downside Risk (Sortino)',
    'cvar': 'Tail Risk (CVaR 5%)',
    'roi': 'Max Return %',
    'winrate': 'Win Rate',
}

# Month helpers for coverage timeline
MONTH_NAMES_SHORT = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun',
                     'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']
_MONTH_IDX = {m: i for i, m in enumerate(MONTH_NAMES_SHORT)}
# Oct-Sep fiscal order (matches the app's display)
_MONTH_ORDER = [9, 10, 11, 0, 1, 2, 3, 4, 5, 6, 7, 8]

# Cell shading colors for coverage timeline
_SHADE_GREEN_HEX = '5E9732'
_SHADE_ORANGE_HEX = 'FF6B35'
_SHADE_GRAY_HEX = 'F0F0F0'


def _interval_to_months(interval_name):
    """Parse 'Sep-Oct' into (8, 9) — 0-indexed month numbers."""
    parts = interval_name.split('-')
    if len(parts) == 2:
        m1 = _MONTH_IDX.get(parts[0])
        m2 = _MONTH_IDX.get(parts[1])
        if m1 is not None and m2 is not None:
            return (m1, m2)
    return None


# =============================================================================
# SHARED HELPERS (mirrors PRF report_generator.py)
# =============================================================================

def _style_table_header_row(table, font_size=Pt(9)):
    """Apply TFC dark-forest header styling to the first row of a table."""
    for cell in table.rows[0].cells:
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), _CLR_DARK_FOREST_HEX)
        shading_elm.set(qn('w:val'), 'clear')
        tcPr.append(shading_elm)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = _CLR_WHITE
                run.font.bold = True
                run.font.size = font_size


def _add_heading_paragraph(doc, text, size, bold=True, color=None,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT,
                           space_before=None, space_after=None):
    """Add a styled heading paragraph (NOT doc.add_heading) for full brand control."""
    para = doc.add_paragraph()
    para.alignment = alignment
    if space_before is not None:
        para.paragraph_format.space_before = space_before
    if space_after is not None:
        para.paragraph_format.space_after = space_after
    run = para.add_run(text)
    run.bold = bold
    run.font.size = size
    if color:
        run.font.color.rgb = color
    return para


def _add_page_field(paragraph):
    """Append a PAGE field code to the given paragraph."""
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run1 = paragraph.add_run()
    run1._element.append(fldChar_begin)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2 = paragraph.add_run()
    run2._element.append(instrText)

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run3 = paragraph.add_run()
    run3._element.append(fldChar_end)


def _add_numpages_field(paragraph):
    """Append a NUMPAGES field code to the given paragraph."""
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run1 = paragraph.add_run()
    run1._element.append(fldChar_begin)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' NUMPAGES '
    run2 = paragraph.add_run()
    run2._element.append(instrText)

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run3 = paragraph.add_run()
    run3._element.append(fldChar_end)


def _add_footer(section):
    """Add branded footer: 'FOR INTERNAL USE ONLY' (left) and 'Page X of Y' (right)."""
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]

    # Add a right-aligned tab stop at the right margin
    pPr = footer_para._element.get_or_add_pPr()
    tabs_elm = OxmlElement('w:tabs')
    tab_elm = OxmlElement('w:tab')
    tab_elm.set(qn('w:val'), 'right')
    tab_elm.set(qn('w:pos'), '13680')  # ~9.5 inches for landscape
    tab_elm.set(qn('w:leader'), 'none')
    tabs_elm.append(tab_elm)
    pPr.append(tabs_elm)

    # Left side
    run_left = footer_para.add_run("FOR INTERNAL USE ONLY")
    run_left.font.bold = True
    run_left.font.size = Pt(8)
    run_left.font.color.rgb = _CLR_MUTED_ROSE

    # Tab to right
    footer_para.add_run("\t")

    # Right side: "Page X of Y"
    run_page = footer_para.add_run("Page ")
    run_page.font.size = Pt(8)
    run_page.font.color.rgb = _CLR_SLATE_BLUE
    _add_page_field(footer_para)
    run_of = footer_para.add_run(" of ")
    run_of.font.size = Pt(8)
    run_of.font.color.rgb = _CLR_SLATE_BLUE
    _add_numpages_field(footer_para)


def _add_simple_table(doc, headers, rows, font_size=Pt(9)):
    """Add a simple table with TFC header styling. Returns the table object."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for idx, col in enumerate(headers):
        hdr_cells[idx].text = col
        for paragraph in hdr_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = font_size
    _style_table_header_row(table, font_size)

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for idx, val in enumerate(row_data):
            row_cells[idx].text = str(val)
            for paragraph in row_cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = font_size

    return table


def _add_chart_to_doc(doc, fig, width=Inches(9.0)):
    """Save a matplotlib figure to a BytesIO buffer and add to document."""
    chart_buffer = BytesIO()
    fig.savefig(chart_buffer, format='png', bbox_inches='tight',
                facecolor=fig.get_facecolor(), dpi=150)
    plt.close(fig)
    chart_buffer.seek(0)
    doc.add_picture(chart_buffer, width=width)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


# =============================================================================
# METRICS HELPERS
# =============================================================================

def _shade_cell(cell, fill_hex, text, font_size=Pt(9), bold=False, text_color=None, center=True):
    """Apply background shading to a cell and set its text."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill_hex)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)

    cell.text = text
    for paragraph in cell.paragraphs:
        if center:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = font_size
            run.font.bold = bold
            if text_color:
                run.font.color.rgb = text_color


def _compute_all_metrics(returns, cost):
    """Compute all portfolio metrics from a returns array and per-acre cost."""
    std = np.std(returns)
    sharpe = float(np.mean(returns) / std) if std > 0 else 0.0
    cvar = float(np.percentile(returns, 5))
    winrate = float(np.mean(returns > 0))
    roi = float((np.mean(returns) + cost) / cost) if cost > 0 else 0.0
    mean_ret = float(np.mean(returns))
    return {
        'sharpe': sharpe, 'cvar': cvar, 'winrate': winrate,
        'roi': roi, 'mean_return': mean_ret, 'cost': cost,
    }


# =============================================================================
# MAIN REPORT GENERATOR
# =============================================================================

def generate_unified_optimizer_report_docx(
    units_data,
    best_combo,
    indep_results,
    metric_key,
    start_year,
    end_year,
    coverage_mode='none',
    coverage_best=None,
    coverage_comparison=None,
    coverage_group_keys=None,
    coverage_metric_name=None,
    get_buyup_intervals_fn=None,
    get_cat_interval_fn=None,
    rate_year=2026,
    stage2_results=None,
    report_stage=1,
):
    """
    Generate a Word document with Unified PRF + AF Joint Optimizer results.

    Args:
        units_data: List of unit dicts from the optimizer (each has
                    yearly_returns, producer_costs, acres, grid_id,
                    type ('PRF' or 'AF'), growing_season (AF), intended_use (PRF),
                    coverage_level, is_cat, productivity, candidates, years, unit_label).
        best_combo: Tuple of best candidate indices per unit.
        indep_results: List of (best_idx, score) tuples from independent optimization.
        metric_key: Optimization metric used ('sharpe', 'cvar', 'roi', 'winrate').
        start_year, end_year: Backtest period.
        coverage_mode: 'none', 'uniform', 'per_category', or 'per_county_crop'.
        coverage_best: Best coverage result (float for uniform, tuple for per_category, etc.).
        coverage_comparison: List of dicts with coverage sweep results.
        coverage_group_keys: Group keys for per_county_crop mode.
        coverage_metric_name: Display name for the coverage metric.
        get_buyup_intervals_fn: Function(growing_season) -> {code: name} dict.
        get_cat_interval_fn: Function(growing_season) -> {code: name} dict.
        rate_year: Rate year for premium calculations.
        stage2_results: Stage 2 HRP rebalancing results dict.

    Returns:
        BytesIO buffer containing the Word document.
    """
    doc = Document()

    # --- Landscape orientation ---
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    # --- Footer ---
    _add_footer(section)

    # =========================================================================
    # COVER PAGE
    # =========================================================================
    doc.add_paragraph()  # blank

    _add_heading_paragraph(doc, "Texas Farm Credit", Pt(26), bold=True,
                           color=_CLR_PRIMARY_GREEN, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    _add_heading_paragraph(doc, "Unified PRF + AF Portfolio", Pt(28), bold=True,
                           color=_CLR_BLACK, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           space_after=Pt(0))
    _add_heading_paragraph(doc, "Optimizer Report", Pt(24), bold=True,
                           color=_CLR_BLACK, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           space_before=Pt(0))
    doc.add_paragraph()

    # Report generated timestamp
    gen_para = doc.add_paragraph()
    gen_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label = gen_para.add_run("Report Generated: ")
    run_label.font.size = Pt(13)
    run_label.font.color.rgb = _CLR_SLATE_BLUE
    run_date = gen_para.add_run(datetime.now().strftime('%Y-%m-%d %H:%M'))
    run_date.font.size = Pt(13)
    run_date.bold = True

    # Generated by
    by_para = doc.add_paragraph()
    by_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_by_label = by_para.add_run("Generated by: ")
    run_by_label.font.size = Pt(13)
    run_by_label.font.color.rgb = _CLR_SLATE_BLUE
    run_by_name = by_para.add_run("Texas Farm Credit")
    run_by_name.font.size = Pt(13)
    run_by_name.bold = True
    run_by_name.font.color.rgb = _CLR_PRIMARY_GREEN

    doc.add_paragraph()

    # Optimization settings summary on cover
    year_count = end_year - start_year + 1
    n_units = len(units_data)
    n_prf = sum(1 for ud in units_data if ud.get('type') == 'PRF')
    n_af = sum(1 for ud in units_data if ud.get('type') != 'PRF')

    settings_lines = [
        f"Optimization Metric: {METRIC_DISPLAY_NAMES.get(metric_key, metric_key)}",
        f"Backtest Period: {start_year} \u2013 {end_year} ({year_count} years)",
        f"Portfolio Units: {n_units} ({n_prf} PRF, {n_af} AF)",
    ]
    if coverage_mode != 'none' and coverage_best is not None:
        if coverage_mode == 'uniform':
            settings_lines.append(f"Optimal Coverage Level: {int(coverage_best * 100)}%")
        elif coverage_mode == 'per_category':
            settings_lines.append(
                f"Optimal Coverage: PRF {int(coverage_best[0] * 100)}% / AF {int(coverage_best[1] * 100)}%"
            )
        elif coverage_mode == 'per_county_crop' and coverage_group_keys:
            parts = []
            for i, key in enumerate(coverage_group_keys):
                if isinstance(coverage_best, (list, tuple)) and i < len(coverage_best):
                    parts.append(f"{key}: {int(coverage_best[i] * 100)}%")
            settings_lines.append("Optimal Coverage: " + ", ".join(parts))

    # Indicate report stage on cover page
    _has_stage2_cover = (stage2_results is not None
                         and 'rebalanced_metrics' in (stage2_results or {}))
    if _has_stage2_cover:
        cover_stage_label = "Stage 2 (Rebalanced)" if report_stage == 2 else "Stage 1 (Original)"
        settings_lines.append(f"Report Version: {cover_stage_label}")

    for line in settings_lines:
        s_para = doc.add_paragraph()
        s_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s_run = s_para.add_run(line)
        s_run.font.size = Pt(12)
        s_run.font.color.rgb = _CLR_SLATE_BLUE

    doc.add_paragraph()

    # Cover disclaimers
    disclaimers = [
        "Past performance is not a guarantee of future returns.",
        "This is a risk management decision-making tool only.",
        f"{rate_year} rates are used for all historical backtesting calculations.",
    ]
    for disc_text in disclaimers:
        disc_para = doc.add_paragraph()
        disc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_disc = disc_para.add_run(disc_text)
        run_disc.italic = True
        run_disc.font.size = Pt(9)
        run_disc.font.color.rgb = _CLR_SLATE_BLUE

    # Beta notice
    doc.add_paragraph()
    beta_para = doc.add_paragraph()
    beta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    beta_run = beta_para.add_run(
        "This report is generated by an application currently in beta. "
        "Results are preliminary and intended for internal evaluation purposes only. "
        "Methodology, outputs, and formatting are subject to change as development continues."
    )
    beta_run.italic = True
    beta_run.font.size = Pt(9)
    beta_run.font.color.rgb = _CLR_MUTED_ROSE

    doc.add_page_break()

    # =========================================================================
    # COMPUTE PORTFOLIO DATA
    # =========================================================================
    total_acres = sum(ud['acres'] for ud in units_data)

    # Joint portfolio returns (acre-weighted, per-acre)
    joint_portfolio = sum(
        units_data[k]['yearly_returns'][best_combo[k]:best_combo[k]+1, :] * units_data[k]['acres']
        for k in range(len(units_data))
    ).flatten() / total_acres

    joint_cost = sum(
        units_data[k]['producer_costs'][best_combo[k]] * units_data[k]['acres']
        for k in range(len(units_data))
    ) / total_acres

    # Independent portfolio returns
    indep_indices = [ir[0] for ir in indep_results]
    indep_portfolio = sum(
        units_data[k]['yearly_returns'][indep_indices[k]:indep_indices[k]+1, :] * units_data[k]['acres']
        for k in range(len(units_data))
    ).flatten() / total_acres

    indep_cost = sum(
        units_data[k]['producer_costs'][indep_indices[k]] * units_data[k]['acres']
        for k in range(len(units_data))
    ) / total_acres

    joint_metrics = _compute_all_metrics(joint_portfolio, joint_cost)
    indep_metrics = _compute_all_metrics(indep_portfolio, indep_cost)

    plot_years = units_data[0]['years']

    # =========================================================================
    # SECTION 1: PORTFOLIO HEADLINE METRICS
    # =========================================================================
    _add_heading_paragraph(doc, "Portfolio Headline Metrics", Pt(20), bold=True,
                           color=_CLR_BLACK, space_before=Pt(6))

    cap = doc.add_paragraph()
    cap_run = cap.add_run(
        "Key performance metrics for the jointly optimized portfolio."
    )
    cap_run.italic = True
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = _CLR_SLATE_BLUE

    # Compute portfolio-level totals for headline display
    total_premium_dollars = sum(
        units_data[k]['producer_costs'][best_combo[k]] * units_data[k]['acres']
        for k in range(len(units_data))
    )
    total_insured_value = sum(
        round(ud.get('cbv', 0) * ud['coverage_level'] * ud['productivity'], 2)
        * ud.get('insurable_interest', 1.0) * ud['acres']
        for ud in units_data
    )

    _has_stage2 = (stage2_results is not None
                   and 'rebalanced_metrics' in (stage2_results or {}))

    # Determine which metrics to display based on report_stage
    if report_stage == 2 and _has_stage2:
        s2m = stage2_results['rebalanced_metrics']
        display_metrics = s2m
        # Recompute Stage 2 total premium from rebalanced acres
        display_total_premium = sum(
            stage2_results['rebalanced_acres'][k] * units_data[k]['producer_costs'][best_combo[k]]
            for k in range(len(units_data))
        )
        display_cost_per_ac = display_total_premium / total_acres if total_acres > 0 else 0.0
        stage_label = "Stage 2 (Rebalanced)"
    else:
        display_metrics = joint_metrics
        display_total_premium = total_premium_dollars
        display_cost_per_ac = joint_cost
        stage_label = "Stage 1 (Original)"

    metrics_headers = ['Metric', 'Value']
    metrics_rows = [
        ['Report Version', stage_label],
        ['Total Insured Value', f"${total_insured_value:,.0f}"],
        ['Total Producer Premium', f"${display_total_premium:,.0f}"],
        ['Total Acres', f"{total_acres:,.0f}"],
        ['', ''],
        ['Sharpe Ratio', f"{display_metrics.get('sharpe', 0):.3f}"],
        ['CVaR (5th Percentile)', f"${display_metrics.get('cvar', 0):,.2f}"],
        ['Win Rate', f"{display_metrics.get('winrate', 0) * 100:.1f}%"],
        ['Return %', f"{display_metrics.get('roi', 0) * 100:.0f}%"],
        ['Mean Net Return/Acre', f"${display_metrics.get('mean_return', 0):,.2f}"],
        ['Producer Premium/Acre', f"${display_cost_per_ac:,.2f}"],
    ]

    # Coverage optimization results
    if coverage_mode != 'none' and coverage_best is not None:
        if coverage_mode == 'uniform':
            metrics_rows.append(['Optimal Coverage Level', f"{int(coverage_best * 100)}%"])
        elif coverage_mode == 'per_category':
            metrics_rows.append(['Optimal Coverage (PRF)', f"{int(coverage_best[0] * 100)}%"])
            metrics_rows.append(['Optimal Coverage (AF)', f"{int(coverage_best[1] * 100)}%"])

    # Stage 2 specific info
    if report_stage == 2 and _has_stage2:
        if stage2_results.get('hrp_enabled'):
            metrics_rows.append(['HRP Rebalancing', 'Enabled'])
            metrics_rows.append(['Turnover Constraint', f"{stage2_results.get('turnover_pct', 10)}%"])
        if stage2_results.get('budget_enabled'):
            metrics_rows.append(['Budget Constraint', f"${stage2_results.get('budget_amount', 0):,.0f}"])
            metrics_rows.append(['Budget Scale Factor', f"{stage2_results.get('budget_scale_factor', 1.0):.4f}"])

    _add_simple_table(doc, metrics_headers, metrics_rows)

    # Check if Stage 2 metrics are materially different from Stage 1
    if _has_stage2:
        s2m_check = stage2_results['rebalanced_metrics']
        s1_sharpe = joint_metrics.get('sharpe', 0)
        s2_sharpe = s2m_check.get('sharpe', 0)
        if abs(s1_sharpe - s2_sharpe) < 0.001:
            note_para = doc.add_paragraph()
            note_run = note_para.add_run(
                "Note: Stage 2 rebalancing produced no material change from Stage 1. "
                "This typically occurs when the turnover constraint is tight or the budget constraint is not binding."
            )
            note_run.italic = True
            note_run.font.size = Pt(8)
            note_run.font.color.rgb = _CLR_SLATE_BLUE

    # =========================================================================
    # SECTION 2: PER-UNIT ALLOCATION TABLE
    # =========================================================================
    doc.add_paragraph()
    _add_heading_paragraph(doc, "Per-Unit Allocation", Pt(20), bold=True,
                           color=_CLR_BLACK, space_before=Pt(6))

    cap2 = doc.add_paragraph()
    cap2_run = cap2.add_run(
        "Optimal interval selection and weight distribution for each portfolio unit."
    )
    cap2_run.italic = True
    cap2_run.font.size = Pt(9)
    cap2_run.font.color.rgb = _CLR_SLATE_BLUE

    alloc_headers = [
        'Unit', 'Type', 'Grid', 'Season', 'Coverage', 'PF', 'II',
        'Acres', 'Selected Intervals', 'Weights',
        'Unit Return %', 'Unit Win Rate'
    ]
    alloc_rows = []

    for k in range(len(units_data)):
        ud = units_data[k]
        cand = ud['candidates'][best_combo[k]]
        unit_type = ud.get('type', 'AF')

        if ud.get('is_cat'):
            if get_cat_interval_fn:
                cat_iv = get_cat_interval_fn(ud['growing_season'])
                cat_name = list(cat_iv.values())[0] if cat_iv else 'Full Season'
            else:
                cat_name = 'Full Season (CAT)'
            selected_names = [f"{cat_name} (CAT)"]
            selected_weights = ['100%']
        elif unit_type == 'PRF':
            # PRF: use INTERVAL_ORDER_11 labels
            selected_names = []
            selected_weights = []
            for idx in range(len(INTERVAL_ORDER_11)):
                if idx < len(cand[1]) and cand[1][idx] > 0.005:
                    selected_names.append(INTERVAL_ORDER_11[idx])
                    selected_weights.append(f"{cand[1][idx] * 100:.0f}%")
        else:
            # AF buy-up
            if get_buyup_intervals_fn:
                intervals_k = get_buyup_intervals_fn(ud['growing_season'])
                codes_k = sorted(intervals_k.keys())
                selected_names = []
                selected_weights = []
                for idx in range(6):
                    if cand[1][idx] > 0:
                        selected_names.append(intervals_k[codes_k[idx]])
                        selected_weights.append(f"{cand[1][idx] * 100:.0f}%")
            else:
                selected_names = [f"Idx {i}" for i in cand[0]]
                selected_weights = [f"{cand[1][i] * 100:.0f}%" for i in cand[0]]

        unit_returns = ud['yearly_returns'][best_combo[k]]
        unit_cost = ud['producer_costs'][best_combo[k]]
        unit_roi = float((np.mean(unit_returns) + unit_cost) / unit_cost) if unit_cost > 0 else 0
        unit_winrate = float(np.mean(unit_returns > 0))

        # Season column: AF shows GS label, PRF shows "PRF"
        if unit_type == 'PRF':
            season_label = 'PRF'
        else:
            gs = ud.get('growing_season', '')
            season_label = f"GS-{gs} ({SEASON_LABELS.get(gs, '')})"

        cov_label = "CAT (65%)" if ud.get('is_cat') else f"{int(ud['coverage_level'] * 100)}%"
        ii_val = ud.get('insurable_interest', 1.0)

        alloc_rows.append([
            ud.get('unit_label', f'Unit {k+1}'),
            unit_type,
            str(ud.get('grid_label') or ud.get('grid_id', '?')),
            season_label,
            cov_label,
            f"{int(ud['productivity'] * 100)}%",
            f"{int(ii_val * 100)}%",
            f"{ud['acres']:,.0f}",
            ', '.join(selected_names),
            ', '.join(selected_weights),
            f"{unit_roi * 100:.0f}%",
            f"{unit_winrate * 100:.1f}%",
        ])

    _add_simple_table(doc, alloc_headers, alloc_rows, font_size=Pt(8))

    # =========================================================================
    # SECTION: GRID ALLOCATION BY INTERVAL
    # =========================================================================
    doc.add_paragraph()
    _add_heading_paragraph(doc, "Grid Allocation by Interval", Pt(20), bold=True,
                           color=_CLR_BLACK, space_before=Pt(6))
    cap_alloc = doc.add_paragraph()
    cap_alloc_run = cap_alloc.add_run(
        "Interval weight distribution across all portfolio units. "
        "Values shown are net estimates \u2014 projected indemnity minus producer premium cost. "
        "Final values are determined by RMA/USDA after interval close."
    )
    cap_alloc_run.italic = True
    cap_alloc_run.font.size = Pt(9)
    cap_alloc_run.font.color.rgb = _CLR_SLATE_BLUE

    prf_units = [(k, ud) for k, ud in enumerate(units_data) if ud.get('type') == 'PRF']
    af_units = [(k, ud) for k, ud in enumerate(units_data) if ud.get('type') != 'PRF']

    if prf_units:
        _add_heading_paragraph(doc, "PRF Units", Pt(14), bold=True,
                               color=_CLR_SLATE_BLUE, space_before=Pt(4), space_after=Pt(2))

        prf_headers = ['Grid'] + INTERVAL_ORDER_11 + ['Acres']
        prf_rows = []
        for k, ud in prf_units:
            cand = ud['candidates'][best_combo[k]]
            grid_label = str(ud.get('grid_label') or ud.get('grid_id', '?'))
            row = [grid_label]
            for idx in range(11):
                if idx < len(cand[1]) and cand[1][idx] > 0.005:
                    row.append(f"{cand[1][idx] * 100:.0f}%")
                else:
                    row.append('-')
            row.append(f"{ud['acres']:,.0f}")
            prf_rows.append(row)

        _add_simple_table(doc, prf_headers, prf_rows, font_size=Pt(8))

    if af_units:
        _add_heading_paragraph(doc, "AF Units", Pt(14), bold=True,
                               color=_CLR_SLATE_BLUE, space_before=Pt(4), space_after=Pt(2))

        # Group AF units by growing season
        af_by_gs = {}
        for k, ud in af_units:
            gs = ud.get('growing_season', 0)
            af_by_gs.setdefault(gs, []).append((k, ud))

        for gs in sorted(af_by_gs.keys()):
            gs_label = SEASON_LABELS.get(gs, f'Season {gs}')
            _add_heading_paragraph(doc, f"Growing Season {gs} \u2014 {gs_label}", Pt(11), bold=True,
                                   color=_CLR_BLACK, space_before=Pt(4), space_after=Pt(2))

            # Get interval names for this GS
            if get_buyup_intervals_fn:
                intervals = get_buyup_intervals_fn(gs)
                codes = sorted(intervals.keys())
                interval_names = [intervals[c] for c in codes]
            else:
                interval_names = [f"Interval {i+1}" for i in range(6)]

            af_headers = ['Grid'] + interval_names + ['Acres']
            af_rows = []

            for k, ud in af_by_gs[gs]:
                cand = ud['candidates'][best_combo[k]]
                grid_label = str(ud.get('grid_label') or ud.get('grid_id', '?'))

                if ud.get('is_cat'):
                    # CAT: all buy-up intervals show '-'
                    row = [f"{grid_label} (CAT)"]
                    for i in range(len(interval_names)):
                        row.append('-')
                    row.append(f"{ud['acres']:,.0f}")
                else:
                    row = [grid_label]
                    for idx in range(6):
                        if idx < len(cand[1]) and cand[1][idx] > 0:
                            row.append(f"{cand[1][idx] * 100:.0f}%")
                        else:
                            row.append('-')
                    row.append(f"{ud['acres']:,.0f}")

                af_rows.append(row)

            _add_simple_table(doc, af_headers, af_rows, font_size=Pt(8))

    # Important disclaimer after allocation tables
    doc.add_paragraph()
    imp_para = doc.add_paragraph()
    imp_run = imp_para.add_run(
        "\u25a0 IMPORTANT: All figures in this report are net estimates \u2014 "
        "projected indemnity minus producer premium cost. They do not represent "
        "guaranteed returns. Final values are determined by RMA/USDA after interval close. "
        "This report is in BETA and should not be used as the sole basis for financial decisions."
    )
    imp_run.bold = True
    imp_run.font.size = Pt(8)
    imp_run.font.color.rgb = _CLR_MUTED_ROSE

    # =========================================================================
    # SECTION 3: PREMIUM & COVERAGE BREAKDOWN
    # =========================================================================
    doc.add_paragraph()
    _add_heading_paragraph(doc, "Premium & Coverage Breakdown", Pt(20), bold=True,
                           color=_CLR_BLACK, space_before=Pt(6))

    # Premium table
    _add_heading_paragraph(doc, "Producer Premium by Unit", Pt(14), bold=True,
                           color=_CLR_SLATE_BLUE, space_before=Pt(4), space_after=Pt(2))

    prem_headers = ['Unit', 'Type', 'Grid', 'Season', 'Coverage', 'Premium/Acre', 'Total Premium']
    prem_rows = []
    total_premium = 0.0

    for k in range(len(units_data)):
        ud = units_data[k]
        unit_cost_per_ac = ud['producer_costs'][best_combo[k]]
        unit_prem_total = unit_cost_per_ac * ud['acres']
        total_premium += unit_prem_total
        unit_type = ud.get('type', 'AF')

        if unit_type == 'PRF':
            season_label = 'PRF'
        else:
            gs = ud.get('growing_season', '')
            season_label = f"GS-{gs} ({SEASON_LABELS.get(gs, '')})"

        cov_label = "CAT (65%)" if ud.get('is_cat') else f"{int(ud['coverage_level'] * 100)}%"

        prem_rows.append([
            ud.get('unit_label', f'Unit {k+1}'),
            unit_type,
            str(ud.get('grid_label') or ud.get('grid_id', '?')),
            season_label,
            cov_label,
            f"${unit_cost_per_ac:,.2f}",
            f"${unit_prem_total:,.0f}",
        ])

    # Total row
    prem_rows.append([
        'TOTAL', '', '', '', '',
        f"${total_premium / total_acres:,.2f}",
        f"${total_premium:,.0f}",
    ])

    _add_simple_table(doc, prem_headers, prem_rows)

    # Coverage table
    doc.add_paragraph()
    _add_heading_paragraph(doc, "Dollar Amount of Protection by Unit", Pt(14), bold=True,
                           color=_CLR_SLATE_BLUE, space_before=Pt(4), space_after=Pt(2))

    cov_headers = ['Unit', 'Type', 'Grid', 'Season', 'CBV', 'II', 'DA Protection/Acre', 'Total Coverage']
    cov_rows = []
    total_coverage = 0.0

    for k in range(len(units_data)):
        ud = units_data[k]
        ii = ud.get('insurable_interest', 1.0)
        cbv_val = ud.get('cbv', 0)
        da_display = round(cbv_val * ud['coverage_level'] * ud['productivity'], 2)
        da_with_ii = da_display * ii
        unit_coverage = da_with_ii * ud['acres']
        total_coverage += unit_coverage
        unit_type = ud.get('type', 'AF')

        if unit_type == 'PRF':
            season_label = 'PRF'
        else:
            gs = ud.get('growing_season', '')
            season_label = f"GS-{gs} ({SEASON_LABELS.get(gs, '')})"

        cov_rows.append([
            ud.get('unit_label', f'Unit {k+1}'),
            unit_type,
            str(ud.get('grid_label') or ud.get('grid_id', '?')),
            season_label,
            f"${cbv_val:,.2f}" if cbv_val else 'N/A',
            f"{int(ii * 100)}%",
            f"${da_with_ii:,.2f}",
            f"${unit_coverage:,.0f}",
        ])

    cov_rows.append([
        'TOTAL', '', '', '', '', '',
        f"${total_coverage / total_acres:,.2f}" if total_acres > 0 else '$0.00',
        f"${total_coverage:,.0f}",
    ])

    _add_simple_table(doc, cov_headers, cov_rows)

    # CAT note if applicable
    has_cat = any(ud.get('is_cat', False) for ud in units_data)
    if has_cat:
        cat_note = doc.add_paragraph()
        cat_run = cat_note.add_run(
            "Note: CAT premium is 100% subsidized ($0/ac rate-based). "
            "Producer pays a $655 administrative fee per crop per county, "
            "not reflected in per-acre calculations."
        )
        cat_run.italic = True
        cat_run.font.size = Pt(8)
        cat_run.font.color.rgb = _CLR_SLATE_BLUE

    # =========================================================================
    # NEW SECTION: COVERAGE LEVEL OPTIMIZATION RESULTS
    # =========================================================================
    if report_stage == 1 and coverage_mode != 'none' and coverage_best is not None:
        doc.add_paragraph()
        _add_heading_paragraph(doc, "Coverage Level Optimization Results", Pt(20), bold=True,
                               color=_CLR_BLACK, space_before=Pt(6))

        # Winning banner
        if coverage_mode == 'uniform':
            banner_text = f"Optimal Coverage Level: {int(coverage_best * 100)}%"
        elif coverage_mode == 'per_category':
            banner_text = (
                f"Optimal Coverage: PRF {int(coverage_best[0] * 100)}% / "
                f"AF {int(coverage_best[1] * 100)}%"
            )
        elif coverage_mode == 'per_county_crop' and coverage_group_keys:
            parts = []
            for i, key in enumerate(coverage_group_keys):
                if isinstance(coverage_best, (list, tuple)) and i < len(coverage_best):
                    parts.append(f"{key}: {int(coverage_best[i] * 100)}%")
            banner_text = "Optimal Coverage: " + ", ".join(parts)
        else:
            banner_text = "Coverage Optimization Complete"

        _add_heading_paragraph(doc, banner_text, Pt(14), bold=True,
                               color=_CLR_PRIMARY_GREEN, space_before=Pt(4), space_after=Pt(4))

        # Comparison table
        if coverage_comparison:
            metric_name = coverage_metric_name or METRIC_DISPLAY_NAMES.get(metric_key, metric_key)
            has_groups = any(entry.get('group') for entry in coverage_comparison)

            cov_comp_headers = ['Best', 'Coverage', metric_name]
            if has_groups:
                cov_comp_headers = ['Best', 'Group', 'Coverage', metric_name]

            # Sort by score descending
            sorted_comparison = sorted(
                coverage_comparison,
                key=lambda e: e.get('score', 0),
                reverse=True
            )

            # Find best score
            best_score = sorted_comparison[0].get('score', 0) if sorted_comparison else None

            cov_comp_rows = []
            for entry in sorted_comparison:
                marker = '\u2b50' if entry.get('score') == best_score else ''
                cov_str = str(entry.get('label', entry.get('coverage', '')))
                score_str = f"{entry.get('score', 0):.2f}"
                if has_groups:
                    cov_comp_rows.append([marker, str(entry.get('group', '')), cov_str, score_str])
                else:
                    cov_comp_rows.append([marker, cov_str, score_str])

            _add_simple_table(doc, cov_comp_headers, cov_comp_rows, font_size=Pt(8))

    # =========================================================================
    # SECTION 5: MONTH COVERAGE TIMELINE
    # =========================================================================
    doc.add_page_break()
    _add_heading_paragraph(doc, "Month Coverage Timeline", Pt(20), bold=True,
                           color=_CLR_BLACK, space_before=Pt(6))

    cap_cov = doc.add_paragraph()
    cap_cov_run = cap_cov.add_run(
        "Monthly coverage map across all portfolio units. "
        "Green indicates single-unit coverage; orange indicates overlap between units."
    )
    cap_cov_run.italic = True
    cap_cov_run.font.size = Pt(9)
    cap_cov_run.font.color.rgb = _CLR_SLATE_BLUE

    # Build coverage map: month_idx -> set of unit indices
    all_months_covered = {}
    for k in range(len(units_data)):
        ud = units_data[k]
        cand = ud['candidates'][best_combo[k]]
        unit_type = ud.get('type', 'AF')

        if ud.get('is_cat'):
            # CAT covers all months spanned by the buy-up intervals
            if get_buyup_intervals_fn:
                for name in get_buyup_intervals_fn(ud['growing_season']).values():
                    months = _interval_to_months(name)
                    if months:
                        for m in months:
                            all_months_covered.setdefault(m, set()).add(k)
        elif unit_type == 'PRF':
            # PRF: iterate INTERVAL_ORDER_11
            for idx in range(len(INTERVAL_ORDER_11)):
                if idx < len(cand[1]) and cand[1][idx] > 0.005:
                    months = _interval_to_months(INTERVAL_ORDER_11[idx])
                    if months:
                        for m in months:
                            all_months_covered.setdefault(m, set()).add(k)
        else:
            # AF buy-up
            if get_buyup_intervals_fn:
                intervals_k = get_buyup_intervals_fn(ud['growing_season'])
                codes_k = sorted(intervals_k.keys())
                for idx in range(6):
                    if cand[1][idx] > 0:
                        name = intervals_k[codes_k[idx]]
                        months = _interval_to_months(name)
                        if months:
                            for m in months:
                                all_months_covered.setdefault(m, set()).add(k)

    month_labels = [MONTH_NAMES_SHORT[m] for m in _MONTH_ORDER]

    # Build the Word table: Unit | Oct | Nov | ... | Sep
    n_cols = 1 + len(_MONTH_ORDER)  # Unit label + 12 months
    cov_table = doc.add_table(rows=1, cols=n_cols)
    cov_table.style = 'Table Grid'
    cov_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row — slate background
    hdr_cells = cov_table.rows[0].cells
    _shade_cell(hdr_cells[0], '5B707F', 'Unit', font_size=Pt(8), bold=True,
                text_color=_CLR_WHITE, center=False)
    for i, m_label in enumerate(month_labels):
        _shade_cell(hdr_cells[i + 1], '5B707F', m_label, font_size=Pt(8),
                    bold=True, text_color=_CLR_WHITE)

    # One row per unit
    for k in range(len(units_data)):
        row_cells = cov_table.add_row().cells
        # Unit label cell
        row_cells[0].text = units_data[k].get('unit_label', f'Unit {k+1}')
        for paragraph in row_cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.bold = True

        # Month cells
        for i, m_idx in enumerate(_MONTH_ORDER):
            units_at_month = all_months_covered.get(m_idx, set())
            if k in units_at_month:
                count = len(units_at_month)
                dots = '\u25cf' * count
                if count >= 2:
                    # Overlap — orange
                    _shade_cell(row_cells[i + 1], _SHADE_ORANGE_HEX, dots,
                                font_size=Pt(8), text_color=_CLR_WHITE)
                else:
                    # Single coverage — green
                    _shade_cell(row_cells[i + 1], _SHADE_GREEN_HEX, dots,
                                font_size=Pt(8), text_color=_CLR_WHITE)
            else:
                # Not covered — gray
                _shade_cell(row_cells[i + 1], _SHADE_GRAY_HEX, '\u2014',
                            font_size=Pt(8))

    # Legend — dot-based, mirrors Streamlit UI
    legend = doc.add_paragraph()
    legend.paragraph_format.space_before = Pt(4)

    legend_items = [
        ('\u25cf', ' = one unit', _CLR_PRIMARY_GREEN),
        ('\u25cf\u25cf', ' = two units', RGBColor(0xFF, 0x6B, 0x35)),
        ('\u25cf\u25cf\u25cf', ' = three units, etc.', RGBColor(0xFF, 0x6B, 0x35)),
    ]
    for dots, label, color in legend_items:
        run_dot = legend.add_run(dots)
        run_dot.font.color.rgb = color
        run_dot.font.size = Pt(8)
        run_lbl = legend.add_run(label + '     ')
        run_lbl.font.size = Pt(7)
        run_lbl.font.color.rgb = RGBColor(0x6B, 0x70, 0x80)  # slate/gray

    run_gray = legend.add_run('\u2014')
    run_gray.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run_gray.font.size = Pt(8)
    run_nc = legend.add_run(' = not covered')
    run_nc.font.size = Pt(7)
    run_nc.font.color.rgb = RGBColor(0x6B, 0x70, 0x80)

    # =========================================================================
    # SECTION 6: HISTORICAL PERFORMANCE CHART
    # =========================================================================
    doc.add_page_break()
    _add_heading_paragraph(doc, "Historical Performance and Analysis", Pt(20), bold=True,
                           color=_CLR_BLACK, space_before=Pt(6))

    cap5 = doc.add_paragraph()
    cap5_run = cap5.add_run(
        "Joint portfolio net return per acre by year. "
        "Years below the break-even line represent net losses (indemnity < premium)."
    )
    cap5_run.italic = True
    cap5_run.font.size = Pt(9)
    cap5_run.font.color.rgb = _CLR_SLATE_BLUE

    try:
        fig, ax = plt.subplots(figsize=(14, 5), dpi=150)
        fig.patch.set_facecolor(_WARM_CREAM)
        ax.set_facecolor(_WARM_CREAM)

        colors = [_HEX_GREEN if v >= 0 else _HEX_ROSE for v in joint_portfolio]
        ax.bar(plot_years.astype(int), joint_portfolio, color=colors, alpha=0.85, width=0.8)

        ax.axhline(y=0, color='black', linewidth=0.8, linestyle='-')
        ax.set_xlabel('Year', fontsize=10)
        ax.set_ylabel('Net Return per Acre ($)', fontsize=10)
        ax.set_title('Joint Portfolio \u2014 Annual Net Return per Acre', fontsize=13, fontweight='bold')

        ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, p: f'${x:,.2f}'))

        # Smart x-axis ticks (show every Nth year based on range)
        all_years = plot_years.astype(int)
        n_years_total = len(all_years)
        tick_step = max(1, n_years_total // 20)
        ax.set_xticks(all_years[::tick_step])
        ax.set_xticklabels(all_years[::tick_step], rotation=45, ha='right', fontsize=7)

        ax.grid(axis='y', alpha=0.3)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

        # Add mean line
        mean_val = float(np.mean(joint_portfolio))
        ax.axhline(y=mean_val, color=_HEX_SLATE, linewidth=1.2, linestyle='--', alpha=0.7)
        ax.annotate(f'Mean: ${mean_val:,.2f}', xy=(all_years[-1], mean_val),
                    fontsize=8, color=_HEX_SLATE, ha='right',
                    xytext=(0, 8), textcoords='offset points')

        plt.tight_layout()
        _add_chart_to_doc(doc, fig)
    except Exception:
        plt.close('all')
        doc.add_paragraph('(Historical performance chart could not be generated.)')

    # --- ROI by Year chart ---
    doc.add_paragraph()
    _add_heading_paragraph(doc, "Annual ROI by Year", Pt(14), bold=True,
                           color=_CLR_SLATE_BLUE, space_before=Pt(4))

    cap_roi = doc.add_paragraph()
    cap_roi_run = cap_roi.add_run(
        "Portfolio ROI (%) by year. ROI = (Indemnity / Premium - 1) \u00d7 100. "
        "Years at -100% received no indemnity (full premium loss)."
    )
    cap_roi_run.italic = True
    cap_roi_run.font.size = Pt(9)
    cap_roi_run.font.color.rgb = _CLR_SLATE_BLUE

    try:
        roi_by_year = (joint_portfolio / joint_cost) * 100 if joint_cost > 0 else np.zeros_like(joint_portfolio)

        fig_roi, ax_roi = plt.subplots(figsize=(14, 5), dpi=150)
        fig_roi.patch.set_facecolor(_WARM_CREAM)
        ax_roi.set_facecolor(_WARM_CREAM)

        roi_colors = [_HEX_GREEN if v >= 0 else _HEX_ROSE for v in roi_by_year]
        ax_roi.bar(all_years, roi_by_year, color=roi_colors, alpha=0.85, width=0.8)

        ax_roi.axhline(y=0, color='black', linewidth=0.8, linestyle='-')
        ax_roi.axhline(y=-100, color='#CC0000', linewidth=1.0, linestyle='--',
                        alpha=0.7, label='Full Premium Loss (-100%)')

        ax_roi.set_xlabel('Year', fontsize=10)
        ax_roi.set_ylabel('ROI (%)', fontsize=10)
        ax_roi.set_title('Joint Portfolio \u2014 Annual ROI (%)', fontsize=13, fontweight='bold')

        ax_roi.yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, p: f'{x:,.0f}%'))

        current_ymin, current_ymax = ax_roi.get_ylim()
        ax_roi.set_ylim(bottom=min(current_ymin, -105), top=current_ymax)

        ax_roi.set_xticks(all_years[::tick_step])
        ax_roi.set_xticklabels(all_years[::tick_step], rotation=45, ha='right', fontsize=7)

        # Mean ROI line
        mean_roi = float(np.mean(roi_by_year))
        ax_roi.axhline(y=mean_roi, color=_HEX_SLATE, linewidth=1.2, linestyle='--', alpha=0.7)
        ax_roi.annotate(f'Mean ROI: {mean_roi:,.0f}%', xy=(all_years[-1], mean_roi),
                        fontsize=8, color=_HEX_SLATE, ha='right',
                        xytext=(0, 8), textcoords='offset points')

        ax_roi.legend(loc='upper left', fontsize=9)
        ax_roi.grid(axis='y', alpha=0.3)
        ax_roi.spines['top'].set_visible(False)
        ax_roi.spines['right'].set_visible(False)

        plt.tight_layout()
        _add_chart_to_doc(doc, fig_roi)
    except Exception:
        plt.close('all')
        doc.add_paragraph('(ROI chart could not be generated.)')

    # =========================================================================
    # SECTION 6b: HRP DENDROGRAM
    # =========================================================================
    if stage2_results and stage2_results.get('hrp_enabled') and len(units_data) >= 3:
        doc.add_paragraph()
        _add_heading_paragraph(doc, "HRP Clustering Dendrogram", Pt(14), bold=True,
                               color=_CLR_SLATE_BLUE, space_before=Pt(4))

        cap_dend = doc.add_paragraph()
        cap_dend_run = cap_dend.add_run(
            "Units joined lower are more correlated. "
            "HRP diversifies across branches."
        )
        cap_dend_run.italic = True
        cap_dend_run.font.size = Pt(9)
        cap_dend_run.font.color.rgb = _CLR_SLATE_BLUE

        try:
            from scipy.cluster.hierarchy import linkage, dendrogram as scipy_dendrogram
            from scipy.spatial.distance import squareform

            # Build labels with type branching
            dend_labels = []
            for k, ud in enumerate(units_data):
                unit_type = ud.get('type', 'AF')
                gl = ud.get('grid_label') or ud.get('grid_id', '?')
                if unit_type == 'AF':
                    dend_labels.append(
                        f"U{k+1}: {SEASON_LABELS.get(ud.get('growing_season', '?'), '?')} \u00b7 {gl}"
                    )
                else:
                    dend_labels.append(
                        f"U{k+1}: PRF \u00b7 {gl}"
                    )

            returns_df = pd.DataFrame({
                dend_labels[k]: units_data[k]['yearly_returns'][best_combo[k]]
                for k in range(len(units_data))
            })
            corr = returns_df.corr().values
            corr = np.clip(corr, -1, 1)
            dist = np.sqrt(0.5 * (1 - corr))
            np.fill_diagonal(dist, 0)
            condensed = squareform(dist, checks=False)
            link = linkage(condensed, method='single')

            fig_dend, ax = plt.subplots(figsize=(14, 4), dpi=150)
            fig_dend.patch.set_facecolor(_WARM_CREAM)
            ax.set_facecolor(_WARM_CREAM)
            scipy_dendrogram(
                link,
                labels=dend_labels,
                leaf_font_size=10,
                color_threshold=0,
                above_threshold_color=_HEX_GREEN,
                ax=ax,
            )
            ax.set_ylabel("Correlation Distance", fontsize=9, color=_HEX_SLATE)
            ax.tick_params(axis='x', labelsize=9)
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            fig_dend.tight_layout()
            _add_chart_to_doc(doc, fig_dend)
            plt.close(fig_dend)
        except Exception:
            plt.close('all')
            doc.add_paragraph('(HRP dendrogram could not be generated.)')

    # =========================================================================
    # SECTION 6: YEAR-BY-YEAR BACKTEST DETAIL TABLE
    # =========================================================================
    doc.add_page_break()
    _add_heading_paragraph(doc, "Year-by-Year Backtest Detail", Pt(20), bold=True,
                           color=_CLR_BLACK, space_before=Pt(6))

    cap6 = doc.add_paragraph()
    cap6_run = cap6.add_run(
        "Per-unit and portfolio-level indemnity, premium, and net return for each historical year. "
        "All values are per-acre."
    )
    cap6_run.italic = True
    cap6_run.font.size = Pt(9)
    cap6_run.font.color.rgb = _CLR_SLATE_BLUE

    # Build portfolio summary and per-unit detail data
    portfolio_summary_rows = []
    # Per-unit data: list of lists, one per unit
    per_unit_rows = [[] for _ in range(len(units_data))]

    for yr_idx, year in enumerate(plot_years):
        port_indem_weighted = 0.0
        port_prem_weighted = 0.0

        for k in range(len(units_data)):
            ud = units_data[k]
            unit_return = ud['yearly_returns'][best_combo[k], yr_idx]
            unit_cost = ud['producer_costs'][best_combo[k]]
            unit_indem = unit_return + unit_cost

            per_unit_rows[k].append([
                str(int(year)),
                f"${unit_indem:,.2f}",
                f"${unit_cost:,.2f}",
                f"${unit_return:,.2f}",
            ])

            port_indem_weighted += unit_indem * ud['acres']
            port_prem_weighted += unit_cost * ud['acres']

        port_indem_ac = port_indem_weighted / total_acres
        port_prem_ac = port_prem_weighted / total_acres
        port_net_ac = port_indem_ac - port_prem_ac
        port_roi = (port_net_ac / port_prem_ac * 100) if port_prem_ac > 0 else 0.0

        portfolio_summary_rows.append([
            str(int(year)),
            f"${port_indem_ac:,.2f}",
            f"${port_prem_ac:,.2f}",
            f"${port_net_ac:,.2f}",
            f"{port_roi:,.0f}%",
        ])

    # Sort descending by year
    portfolio_summary_rows.sort(key=lambda r: int(r[0]), reverse=True)
    for k in range(len(units_data)):
        per_unit_rows[k].sort(key=lambda r: int(r[0]), reverse=True)

    # Summary stats row at top
    best_yr_idx = int(np.argmax(joint_portfolio))
    worst_yr_idx = int(np.argmin(joint_portfolio))
    best_yr = int(plot_years[best_yr_idx])
    worst_yr = int(plot_years[worst_yr_idx])

    # Cumulative ROI = total indemnity / total premium across all years
    total_indem_all_years = float(np.sum(joint_portfolio)) + joint_cost * len(plot_years)
    total_prem_all_years = joint_cost * len(plot_years)
    cumulative_roi = (total_indem_all_years / total_prem_all_years * 100) if total_prem_all_years > 0 else 0

    summary_para = doc.add_paragraph()
    summary_parts = [
        f"Years: {len(plot_years)}",
        f"Avg Net/ac: ${np.mean(joint_portfolio):,.2f}",
        f"Cumulative ROI: {cumulative_roi:,.0f}%",
        f"Win Rate: {np.mean(joint_portfolio > 0) * 100:.1f}%",
        f"Best Year: {best_yr} (${np.max(joint_portfolio):,.2f})",
        f"Worst Year: {worst_yr} (${np.min(joint_portfolio):,.2f})",
    ]
    summary_run = summary_para.add_run("  |  ".join(summary_parts))
    summary_run.font.size = Pt(9)
    summary_run.font.color.rgb = _CLR_SLATE_BLUE
    summary_run.bold = True

    # Portfolio Summary Table (per-acre)
    _add_heading_paragraph(doc, "Portfolio Summary (Per Acre)", Pt(14), bold=True,
                           color=_CLR_SLATE_BLUE, space_before=Pt(4), space_after=Pt(2))
    _add_simple_table(doc, ['Year', 'Portfolio Indemnity/ac', 'Portfolio Premium/ac',
                            'Portfolio Net/ac', 'Portfolio ROI'],
                      portfolio_summary_rows, font_size=Pt(8))

    # Per-Unit Detail Tables (per-acre)
    if len(units_data) >= 4:
        doc.add_page_break()
    _add_heading_paragraph(doc, "Per-Unit Detail (Per Acre)", Pt(14), bold=True,
                           color=_CLR_SLATE_BLUE, space_before=Pt(6), space_after=Pt(2))

    for k in range(len(units_data)):
        ud = units_data[k]
        unit_type = ud.get('type', 'AF')
        grid_label = ud.get('grid_label') or ud.get('grid_id', '?')
        unit_label = ud.get('unit_label', f'Unit {k+1}')
        _add_heading_paragraph(doc, f"{unit_label}: {unit_type} Grid {grid_label}", Pt(11), bold=True,
                               color=_CLR_BLACK, space_before=Pt(4), space_after=Pt(2))
        _add_simple_table(doc, ['Year', 'Indemnity/ac', 'Premium/ac', 'Net/ac'],
                          per_unit_rows[k], font_size=Pt(8))

    # --- Gross Dollar Year-by-Year Table ---
    doc.add_paragraph()
    _add_heading_paragraph(doc, "Year-by-Year Backtest Detail (Gross Dollars)", Pt(20), bold=True,
                           color=_CLR_BLACK, space_before=Pt(6))

    cap6g = doc.add_paragraph()
    cap6g_run = cap6g.add_run(
        "Total dollar amounts (per-acre values scaled by each unit's insured acres). "
        "Portfolio columns reflect the aggregate across all units."
    )
    cap6g_run.italic = True
    cap6g_run.font.size = Pt(9)
    cap6g_run.font.color.rgb = _CLR_SLATE_BLUE

    gross_portfolio_rows = []
    gross_per_unit_rows = [[] for _ in range(len(units_data))]
    gross_total_indem = 0.0
    gross_total_prem = 0.0

    for yr_idx, year in enumerate(plot_years):
        port_indem_gross = 0.0
        port_prem_gross = 0.0

        for k in range(len(units_data)):
            ud = units_data[k]
            unit_return_ac = ud['yearly_returns'][best_combo[k], yr_idx]
            unit_cost_ac = ud['producer_costs'][best_combo[k]]
            unit_indem_ac = unit_return_ac + unit_cost_ac

            unit_indem_gross = unit_indem_ac * ud['acres']
            unit_prem_gross = unit_cost_ac * ud['acres']
            unit_net_gross = unit_indem_gross - unit_prem_gross

            gross_per_unit_rows[k].append([
                str(int(year)),
                f"${unit_indem_gross:,.0f}",
                f"${unit_prem_gross:,.0f}",
                f"${unit_net_gross:,.0f}",
            ])

            port_indem_gross += unit_indem_gross
            port_prem_gross += unit_prem_gross

        port_net_gross = port_indem_gross - port_prem_gross
        port_roi = (port_net_gross / port_prem_gross * 100) if port_prem_gross > 0 else 0.0

        gross_portfolio_rows.append([
            str(int(year)),
            f"${port_indem_gross:,.0f}",
            f"${port_prem_gross:,.0f}",
            f"${port_net_gross:,.0f}",
            f"{port_roi:,.0f}%",
        ])

        gross_total_indem += port_indem_gross
        gross_total_prem += port_prem_gross

    gross_portfolio_rows.sort(key=lambda r: int(r[0]), reverse=True)
    for k in range(len(units_data)):
        gross_per_unit_rows[k].sort(key=lambda r: int(r[0]), reverse=True)

    # Gross summary stats
    gross_cum_roi = (gross_total_indem / gross_total_prem * 100) if gross_total_prem > 0 else 0
    gross_avg_net = (gross_total_indem - gross_total_prem) / len(plot_years)

    gross_summary = doc.add_paragraph()
    gross_parts = [
        f"Years: {len(plot_years)}",
        f"Avg Net: ${gross_avg_net:,.0f}",
        f"Cumulative ROI: {gross_cum_roi:,.0f}%",
        f"Total Premium ({len(plot_years)} yr): ${gross_total_prem:,.0f}",
        f"Total Indemnity ({len(plot_years)} yr): ${gross_total_indem:,.0f}",
    ]
    gross_summary_run = gross_summary.add_run("  |  ".join(gross_parts))
    gross_summary_run.font.size = Pt(9)
    gross_summary_run.font.color.rgb = _CLR_SLATE_BLUE
    gross_summary_run.bold = True

    # Portfolio Summary Table (gross dollars)
    _add_heading_paragraph(doc, "Portfolio Summary (Gross Dollars)", Pt(14), bold=True,
                           color=_CLR_SLATE_BLUE, space_before=Pt(4), space_after=Pt(2))
    _add_simple_table(doc, ['Year', 'Portfolio Indemnity', 'Portfolio Premium',
                            'Portfolio Net', 'Portfolio ROI'],
                      gross_portfolio_rows, font_size=Pt(8))

    # Per-Unit Detail Tables (gross dollars)
    if len(units_data) >= 4:
        doc.add_page_break()
    _add_heading_paragraph(doc, "Per-Unit Detail (Gross Dollars)", Pt(14), bold=True,
                           color=_CLR_SLATE_BLUE, space_before=Pt(6), space_after=Pt(2))

    for k in range(len(units_data)):
        ud = units_data[k]
        unit_type = ud.get('type', 'AF')
        grid_label = ud.get('grid_label') or ud.get('grid_id', '?')
        unit_label = ud.get('unit_label', f'Unit {k+1}')
        _add_heading_paragraph(doc, f"{unit_label}: {unit_type} Grid {grid_label}", Pt(11), bold=True,
                               color=_CLR_BLACK, space_before=Pt(4), space_after=Pt(2))
        _add_simple_table(doc, ['Year', 'Indemnity', 'Premium', 'Net'],
                          gross_per_unit_rows[k], font_size=Pt(8))

    # =========================================================================
    # SECTION 7: OPTIMIZATION SETTINGS
    # =========================================================================
    doc.add_page_break()
    _add_heading_paragraph(doc, "Optimization Settings & Configuration", Pt(20), bold=True,
                           color=_CLR_BLACK, space_before=Pt(6))

    settings = doc.add_paragraph()

    settings.add_run("Optimization Metric: ").bold = True
    settings.add_run(f"{METRIC_DISPLAY_NAMES.get(metric_key, metric_key)}\n")

    settings.add_run("Backtest Period: ").bold = True
    settings.add_run(f"{start_year} \u2013 {end_year} ({year_count} years)\n")

    settings.add_run("Number of Units: ").bold = True
    settings.add_run(f"{n_units} ({n_prf} PRF, {n_af} AF)\n")

    settings.add_run("Total Acres: ").bold = True
    settings.add_run(f"{total_acres:,.0f}\n")

    # Coverage mode reporting
    if coverage_mode != 'none' and coverage_best is not None:
        settings.add_run("Coverage Level Optimization: ").bold = True
        if coverage_mode == 'uniform':
            settings.add_run(f"Enabled (Uniform) \u2014 Best: {int(coverage_best * 100)}%\n")
        elif coverage_mode == 'per_category':
            settings.add_run(
                f"Enabled (Per-Category) \u2014 PRF: {int(coverage_best[0] * 100)}%, "
                f"AF: {int(coverage_best[1] * 100)}%\n"
            )
        elif coverage_mode == 'per_county_crop':
            parts = []
            for i, key in enumerate(coverage_group_keys or []):
                if isinstance(coverage_best, (list, tuple)) and i < len(coverage_best):
                    parts.append(f"{key}: {int(coverage_best[i] * 100)}%")
            settings.add_run(f"Enabled (Per-County/Crop) \u2014 {', '.join(parts)}\n")
    else:
        settings.add_run("Coverage Level Optimization: ").bold = True
        settings.add_run("Disabled (per-unit selection)\n")

    settings.add_run("\nPer-Unit Configuration:\n").bold = True

    for k in range(len(units_data)):
        ud = units_data[k]
        unit_type = ud.get('type', 'AF')
        cov_label = "CAT (65%)" if ud.get('is_cat') else f"{int(ud['coverage_level'] * 100)}%"
        ii_val = ud.get('insurable_interest', 1.0)
        ii_str = f"II {int(ii_val * 100)}%, " if ii_val < 1.0 else ""
        grid_id = ud.get('grid_label') or ud.get('grid_id', '?')

        if unit_type == 'PRF':
            # Build intervals string for PRF
            cand = ud['candidates'][best_combo[k]]
            iv_parts = []
            for idx in range(len(INTERVAL_ORDER_11)):
                if idx < len(cand[1]) and cand[1][idx] > 0.005:
                    iv_parts.append(f"{INTERVAL_ORDER_11[idx]} ({cand[1][idx]*100:.0f}%)")
            intervals_str = ', '.join(iv_parts) if iv_parts else 'None'

            settings.add_run(f"  Unit {k+1} (PRF): ").bold = True
            settings.add_run(
                f"Grid {grid_id}, PRF, "
                f"Cov {cov_label}, "
                f"PF {int(ud['productivity'] * 100)}%, "
                f"{ii_str}"
                f"{ud['acres']:,.0f} ac \u00b7 {intervals_str}\n"
            )
        else:
            gs = ud.get('growing_season', '')
            gs_label = f"GS-{gs} ({SEASON_LABELS.get(gs, '')})"

            # Build intervals string for AF
            cand = ud['candidates'][best_combo[k]]
            if ud.get('is_cat'):
                if get_cat_interval_fn:
                    cat_iv = get_cat_interval_fn(ud['growing_season'])
                    cat_name = list(cat_iv.values())[0] if cat_iv else 'Full Season'
                else:
                    cat_name = 'Full Season'
                intervals_str = f"{cat_name} (CAT 100%)"
            elif get_buyup_intervals_fn:
                intervals_k = get_buyup_intervals_fn(ud['growing_season'])
                codes_k = sorted(intervals_k.keys())
                iv_parts = []
                for idx in range(6):
                    if cand[1][idx] > 0:
                        iv_parts.append(f"{intervals_k[codes_k[idx]]} ({cand[1][idx]*100:.0f}%)")
                intervals_str = ', '.join(iv_parts) if iv_parts else 'None'
            else:
                intervals_str = 'N/A'

            settings.add_run(f"  {ud.get('unit_label', f'Unit {k+1}')} (AF): ").bold = True
            settings.add_run(
                f"Grid {grid_id}, {gs_label}, "
                f"Coverage {cov_label}, "
                f"PF {int(ud['productivity'] * 100)}%, "
                f"{ii_str}"
                f"{ud['acres']:,.0f} ac \u00b7 {intervals_str}\n"
            )

    # Methodology note
    settings.add_run("\nMethodology: ").bold = True
    if n_units <= 3:
        settings.add_run(
            f"Exhaustive {'2' if n_units == 2 else '3'}-unit joint optimization. "
            f"All valid interval/weight combinations were tested across all units simultaneously.\n"
        )
    else:
        settings.add_run(
            f"Greedy sequential pairing heuristic ({n_units} units). "
            f"Results are highly optimized but not guaranteed to be the global maximum.\n"
        )

    # =========================================================================
    # GLOSSARY
    # =========================================================================
    doc.add_page_break()
    _add_heading_paragraph(doc, "Glossary", Pt(20), bold=True,
                           color=_CLR_BLACK, space_before=Pt(6))

    glossary_entries = [
        ('Annual Forage (AF)',
         'A federally-subsidized rainfall index insurance program for annually planted crops '
         'used for livestock feed. Commodity Code 0332. Covers 12 growing seasons with '
         '6 buy-up intervals and 1 CAT interval per season.'),
        ('Buy-Up Coverage',
         'Standard coverage at levels from 70% to 90%. The producer selects up to 3 '
         'non-adjacent two-month intervals and assigns percentage weights summing to 100%.'),
        ('CAT (Catastrophic) Coverage',
         'Minimum coverage at the 65% level using a single full-season interval at 100% weight. '
         'Premium is fully subsidized; the producer pays a $655 administrative fee per crop per county.'),
        ('County Base Value (CBV)',
         'The dollar value per acre established by USDA for a given grid, growing season, and crop year. '
         'Serves as the foundation for all protection and premium calculations.'),
        ('Coverage Level',
         'The percentage of the expected index (100) at which an indemnity is triggered. '
         'At 90% coverage, a payout occurs when the actual rainfall index falls below 90.'),
        ('CVaR (Conditional Value-at-Risk)',
         'The average loss in the worst 5% of historical outcomes. A tail-risk measure that '
         'answers: "How bad does it get in the worst-case scenarios?" More negative = worse.'),
        ('Dollar Amount of Protection',
         'County Base Value x Coverage Level x Productivity Factor. The per-acre insured value '
         'before applying interval weights and acres.'),
        ('Growing Season',
         'One of 12 planting windows (GS-1 through GS-12), each associated with a set of '
         '6 consecutive two-month rainfall intervals. Named by planting month (e.g., GS-1 = August).'),
        ('Indemnity',
         'The insurance payout triggered when the actual rainfall index falls below the coverage level trigger. '
         'Calculated as: (1 - Actual Index / Trigger) x Policy Protection per interval.'),
        ('Insurable Interest',
         'The percentage of financial risk borne by the producer (0-100%). Scales coverage, '
         'premium, and indemnity linearly.'),
        ('Intended Use',
         'The purpose for which the insured forage is produced. PRF offers Grazing (code 7) '
         'and Haying (code 30), each with different premium rates and allocation bounds.'),
        ('Interval',
         'A two-month rainfall measurement period (e.g., Sep-Oct, Oct-Nov). '
         'Each growing season has 6 buy-up intervals. Adjacent intervals share a month '
         'and cannot both be selected within the same growing season.'),
        ('Irrigation Practice',
         'Applies to PRF Haying only. Options are Irrigated and Non-Irrigated, '
         'affecting premium rates.'),
        ('Joint Optimization',
         'Portfolio-level optimization that considers the correlation of returns across '
         'multiple units simultaneously, rather than optimizing each unit independently. '
         'Reduces portfolio risk through diversification.'),
        ('Net Return',
         'Indemnity minus producer premium for a given year. Positive = profitable year; '
         'negative = premium exceeded payouts.'),
        ('Organic Practice',
         'Applies to PRF Haying only. Options include No Organic Practice Specified, '
         'Organic (Transitional Acreage), and Organic (100% Organic Acreage).'),
        ('Pasture, Rangeland, Forage (PRF)',
         'A federally-subsidized rainfall index insurance program for perennial forage used '
         'for livestock grazing or haying. Uses 11 two-month rainfall intervals (Jan-Feb '
         'through Nov-Dec) with 2-6 intervals selected per policy.'),
        ('Producer Premium',
         'The out-of-pocket cost to the producer after the federal subsidy is applied. '
         'Calculated as: Total Premium - Premium Subsidy.'),
        ('Productivity Factor (PF)',
         'A scalar (60-150%) that adjusts the Dollar Amount of Protection to reflect '
         'the productivity of the insured acreage relative to the county average.'),
        ('ROI (Return on Investment)',
         'Total indemnity divided by total producer premium, expressed as a percentage. '
         'ROI of 100% means the producer broke even; above 100% is profitable.'),
        ('Sharpe Ratio',
         'Mean net return divided by the standard deviation of returns. Measures '
         'risk-adjusted performance. Higher = better return per unit of risk.'),
        ('Subsidy',
         'The federal government\'s share of the total premium. Varies by coverage level '
         '(higher coverage = lower subsidy percentage).'),
        ('Unified Portfolio',
         'A portfolio combining PRF and AF insurance units across different programs, grids, '
         'and growing seasons. Joint optimization exploits cross-program diversification '
         'to improve risk-adjusted returns.'),
        ('Win Rate',
         'The percentage of historical years where the portfolio generated a positive net return '
         '(indemnity exceeded producer premium).'),
    ]

    glossary_headers = ['Term', 'Definition']
    glossary_rows = [[term, defn] for term, defn in glossary_entries]
    _add_simple_table(doc, glossary_headers, glossary_rows, font_size=Pt(8))

    # =========================================================================
    # DISCLAIMER
    # =========================================================================
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.paragraph_format.space_before = Pt(12)
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disc_texts = [
        "Past performance is not a guarantee of future returns.",
        "This is a risk management decision-making tool only.",
        f"{rate_year} rates are used for all historical backtesting calculations.",
    ]
    for i, txt in enumerate(disc_texts):
        run = disclaimer.add_run(txt + ("\n" if i < len(disc_texts) - 1 else ""))
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = _CLR_SLATE_BLUE

    # =========================================================================
    # SAVE & RETURN
    # =========================================================================
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
